"""
Resume processing module for the AI Job Qualification Screening System.

This module handles resume file uploads, text extraction, and AI-powered
resume analysis for enhanced job qualification matching.
"""

import os
import hashlib
import json
import logging
from datetime import datetime
from typing import Dict, Optional, Any
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF processing will be disabled.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be disabled.")

from src.utils.logger import JobAutomationLogger

logger = JobAutomationLogger()


class ResumeProcessor:
    """
    Handles resume file processing including text extraction and AI analysis.
    
    This class provides functionality to:
    - Extract text from PDF and DOCX files
    - Process resume text with AI to extract structured information
    - Calculate file hashes for duplicate detection
    """
    
    def __init__(self, ai_client=None):
        """
        Initialize the resume processor.
        
        Args:
            ai_client: AI client for resume analysis (optional)
        """
        self.ai_client = ai_client
        self.supported_extensions = {'.pdf', '.docx'}
        
        if not PDF_AVAILABLE:
            self.supported_extensions.discard('.pdf')
        if not DOCX_AVAILABLE:
            self.supported_extensions.discard('.docx')
    
    def is_supported_file(self, filename: str) -> bool:
        """
        Check if the file type is supported for processing.
        
        Args:
            filename: Name of the file to check
            
        Returns:
            True if the file type is supported, False otherwise
        """
        if not filename:
            return False
        
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_extensions
    
    def calculate_file_hash(self, file_path: str) -> str:
        """
        Calculate SHA256 hash of a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            SHA256 hash of the file content
        """
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
                return hashlib.sha256(file_content).hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash for {file_path}: {e}")
            raise
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text from the PDF
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is not available. Cannot process PDF files.")
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the DOCX
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not available. Cannot process DOCX files.")
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from a resume file (PDF or DOCX).
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text from the file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def process_resume_with_ai(self, file_path: str) -> Dict[str, Any]:
        """
        Process resume text with AI and return structured data.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Structured resume data extracted by AI
        """
        if not self.ai_client:
            logger.warning("No AI client available. Returning basic text extraction.")
            return self._extract_basic_resume_data(file_path)
        
        try:
            # Extract text from file
            resume_text = self.extract_text_from_file(file_path)
            
            # Create AI prompt for structured extraction
            prompt = self._create_resume_analysis_prompt(resume_text)
            
            # Call AI API
            response = self.ai_client.analyze_text(prompt)
            
            # Parse AI response
            try:
                processed_data = json.loads(response)
                logger.info(f"Successfully processed resume with AI: {file_path}")
                return processed_data
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse AI response as JSON: {e}")
                logger.error(f"AI Response: {response}")
                # Fall back to basic extraction
                return self._extract_basic_resume_data(file_path)
                
        except Exception as e:
            logger.error(f"Error processing resume with AI {file_path}: {e}")
            # Fall back to basic extraction
            return self._extract_basic_resume_data(file_path)
    
    def _create_resume_analysis_prompt(self, resume_text: str) -> str:
        """
        Create the AI prompt for resume analysis.
        
        Args:
            resume_text: Extracted text from the resume
            
        Returns:
            Formatted prompt for AI analysis
        """
        return f"""
        Extract structured information from this resume and return as JSON.
        
        Resume Text:
        {resume_text}
        
        Please extract and return the following information in JSON format:
        {{
            "personal_info": {{
                "name": "",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin": ""
            }},
            "education": [
                {{
                    "degree": "",
                    "field": "",
                    "school": "",
                    "graduation_year": "",
                    "gpa": ""
                }}
            ],
            "experience": [
                {{
                    "title": "",
                    "company": "",
                    "duration": "",
                    "description": "",
                    "key_achievements": []
                }}
            ],
            "skills": {{
                "technical": [],
                "programming_languages": [],
                "frameworks": [],
                "tools": [],
                "soft_skills": []
            }},
            "certifications": [],
            "projects": [
                {{
                    "name": "",
                    "description": "",
                    "technologies": [],
                    "url": ""
                }}
            ],
            "total_years_experience": 0,
            "summary": ""
        }}
        
        Important: Return ONLY valid JSON. Do not include any explanatory text.
        """
    
    def _extract_basic_resume_data(self, file_path: str) -> Dict[str, Any]:
        """
        Extract basic resume data when AI processing is not available.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Basic structured resume data
        """
        try:
            text = self.extract_text_from_file(file_path)
            
            # Basic extraction logic
            lines = text.split('\n')
            skills = []
            experience = []
            
            # Simple keyword-based extraction
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in ['python', 'java', 'javascript', 'react', 'node.js', 'aws', 'docker']):
                    skills.append(line.strip())
                if any(keyword in line_lower for keyword in ['experience', 'work', 'job', 'position']):
                    experience.append(line.strip())
            
            return {
                "personal_info": {
                    "name": "",
                    "email": "",
                    "phone": "",
                    "location": "",
                    "linkedin": ""
                },
                "education": [],
                "experience": [{"title": "Extracted from resume", "description": exp} for exp in experience[:3]],
                "skills": {
                    "technical": skills,
                    "programming_languages": [],
                    "frameworks": [],
                    "tools": [],
                    "soft_skills": []
                },
                "certifications": [],
                "projects": [],
                "total_years_experience": 0,
                "summary": text[:500] + "..." if len(text) > 500 else text
            }
            
        except Exception as e:
            logger.error(f"Error in basic resume extraction {file_path}: {e}")
            return {
                "personal_info": {},
                "education": [],
                "experience": [],
                "skills": {"technical": [], "programming_languages": [], "frameworks": [], "tools": [], "soft_skills": []},
                "certifications": [],
                "projects": [],
                "total_years_experience": 0,
                "summary": "Error extracting resume data"
            }
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Get metadata about a resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing file metadata
        """
        try:
            stat = os.stat(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            return {
                "file_size": stat.st_size,
                "file_type": file_extension[1:] if file_extension else "",
                "created_time": datetime.fromtimestamp(stat.st_ctime),
                "modified_time": datetime.fromtimestamp(stat.st_mtime),
                "is_supported": self.is_supported_file(file_path)
            }
        except Exception as e:
            logger.error(f"Error getting file metadata for {file_path}: {e}")
            return {} 